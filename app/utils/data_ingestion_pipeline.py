"""
Data Ingestion Pipeline for Dynamic RAG System.

What: Extracts content from PDFs/TXT/DOCX/MD, normalizes tables and OCRs
      images (graphs) into text, chunks into ~512-token segments with overlap,
      and attaches metadata.

Why:  Convert heterogeneous file inputs into semantically searchable text
      chunks suitable for embedding and retrieval.

How:  Uses pdfplumber for page text, Camelot for tables, pytesseract for OCR,
      and NLTK sentence/word tokenizers for chunking with configurable size.
"""

import io
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import re

# PDF processing
import pdfplumber
import camelot
from camelot.core import TableList

# Document processing
import docx
from docx import Document

# Image processing for graphs
import pytesseract
from PIL import Image
import fitz  # PyMuPDF

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

from app.config import FILE_PROCESSING_CONFIG

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download required NLTK data (handle newer NLTK splitting 'punkt' resources)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt')
    except Exception:
        pass

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    try:
        nltk.download('punkt_tab')
    except Exception:
        pass


@dataclass
class Document:
    """Document structure for processed content."""
    content: str
    metadata: Dict[str, Any]
    page_number: Optional[int] = None
    content_type: str = "text"  # text, table, graph


class DataIngestionPipeline:
    """Pipeline for extracting and processing content from various file formats."""
    
    def __init__(self):
        """Initialize chunking configuration from settings."""
        self.chunk_size = FILE_PROCESSING_CONFIG["chunk_size"]
        self.chunk_overlap = FILE_PROCESSING_CONFIG["chunk_overlap"]
        
    async def process_files(self, files: List[Tuple[str, bytes]]) -> List[Document]:
        """Process multiple files and return normalized `Document` chunks.

        Args:
            files: Sequence of (filename, raw_bytes).

        Returns:
            Flattened list of chunked Documents across all inputs.
        """
        all_documents = []
        
        for filename, file_content in files:
            try:
                logger.info(f"Processing file: {filename}")
                documents = await self._process_single_file(filename, file_content)
                all_documents.extend(documents)
                logger.info(f"Extracted {len(documents)} documents from {filename}")
                
            except Exception as e:
                logger.error(f"Error processing {filename}: {str(e)}")
                # Continue processing other files
                continue
                
        return all_documents
    
    async def _process_single_file(self, filename: str, file_content: bytes) -> List[Document]:
        """Dispatch file processing based on extension (pdf/txt/docx/md)."""
        file_path = Path(filename)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return await self._process_pdf(filename, file_content)
        elif extension == '.txt':
            return await self._process_text(filename, file_content)
        elif extension == '.docx':
            return await self._process_docx(filename, file_content)
        elif extension == '.md':
            return await self._process_markdown(filename, file_content)
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return []
    
    async def _process_pdf(self, filename: str, file_content: bytes) -> List[Document]:
        """Extract text (pdfplumber), tables (Camelot), and OCR captions from a PDF."""
        documents = []
        
        try:
            # Create a BytesIO object for pdfplumber
            pdf_io = io.BytesIO(file_content)
            
            with pdfplumber.open(pdf_io) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text and text.strip():
                        text_docs = self._chunk_text(
                            text, 
                            metadata={
                                "filename": filename,
                                "page_number": page_num,
                                "type": "text"
                            }
                        )
                        documents.extend(text_docs)
                    
                    # Extract tables using camelot
                    try:
                        tables = camelot.read_pdf(
                            pdf_io, 
                            pages=str(page_num),
                            flavor='lattice'
                        )
                        
                        for table_idx, table in enumerate(tables):
                            table_text = self._table_to_text(table.df)
                            if table_text:
                                table_docs = self._chunk_text(
                                    table_text,
                                    metadata={
                                        "filename": filename,
                                        "page_number": page_num,
                                        "type": "table",
                                        "table_index": table_idx
                                    }
                                )
                                documents.extend(table_docs)
                                
                    except Exception as e:
                        logger.warning(f"Table extraction failed for page {page_num}: {str(e)}")
                    
                    # Extract images/graphs
                    try:
                        images = page.images
                        for img_idx, img in enumerate(images):
                            # Convert image to text using OCR
                            img_text = await self._extract_image_text(img, page)
                            if img_text:
                                graph_docs = self._chunk_text(
                                    img_text,
                                    metadata={
                                        "filename": filename,
                                        "page_number": page_num,
                                        "type": "graph",
                                        "image_index": img_idx
                                    }
                                )
                                documents.extend(graph_docs)
                                
                    except Exception as e:
                        logger.warning(f"Image extraction failed for page {page_num}: {str(e)}")
                        
        except Exception as e:
            logger.error(f"PDF processing failed for {filename}: {str(e)}")
            
        return documents
    
    async def _process_text(self, filename: str, file_content: bytes) -> List[Document]:
        """Decode and chunk a plain text file into sentence-based segments."""
        try:
            text = file_content.decode('utf-8')
            return self._chunk_text(
                text,
                metadata={
                    "filename": filename,
                    "type": "text"
                }
            )
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    return self._chunk_text(
                        text,
                        metadata={
                            "filename": filename,
                            "type": "text"
                        }
                    )
                except UnicodeDecodeError:
                    continue
            logger.error(f"Could not decode text file: {filename}")
            return []
    
    async def _process_docx(self, filename: str, file_content: bytes) -> List[Document]:
        """Extract paragraphs and tables from DOCX and create chunks."""
        documents = []
        
        try:
            doc_io = io.BytesIO(file_content)
            doc = Document(doc_io)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if text_content:
                full_text = "\n".join(text_content)
                text_docs = self._chunk_text(
                    full_text,
                    metadata={
                        "filename": filename,
                        "type": "text"
                    }
                )
                documents.extend(text_docs)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = self._table_to_text(table)
                if table_text:
                    table_docs = self._chunk_text(
                        table_text,
                        metadata={
                            "filename": filename,
                            "type": "table",
                            "table_index": table_idx
                        }
                    )
                    documents.extend(table_docs)
                    
        except Exception as e:
            logger.error(f"DOCX processing failed for {filename}: {str(e)}")
            
        return documents
    
    async def _process_markdown(self, filename: str, file_content: bytes) -> List[Document]:
        """Strip basic Markdown syntax and chunk the resulting text."""
        try:
            text = file_content.decode('utf-8')
            # Simple markdown processing - remove markdown syntax
            clean_text = re.sub(r'#{1,6}\s+', '', text)  # Remove headers
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)  # Remove bold
            clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)  # Remove italic
            clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)  # Remove code
            
            return self._chunk_text(
                clean_text,
                metadata={
                    "filename": filename,
                    "type": "text"
                }
            )
        except Exception as e:
            logger.error(f"Markdown processing failed for {filename}: {str(e)}")
            return []
    
    def _table_to_text(self, table) -> str:
        """Convert a Camelot/docx table into pipe-delimited row text."""
        try:
            if hasattr(table, 'df'):  # Camelot table
                df = table.df
            else:  # docx table
                df = table
            
            # Convert to text representation
            rows = []
            for _, row in df.iterrows() if hasattr(df, 'iterrows') else enumerate(df.rows):
                if hasattr(row, 'tolist'):
                    row_text = " | ".join(str(cell) for cell in row.tolist() if cell)
                else:
                    row_text = " | ".join(str(cell.text) for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    rows.append(row_text)
            
            return "\n".join(rows)
            
        except Exception as e:
            logger.warning(f"Table conversion failed: {str(e)}")
            return ""
    
    async def _extract_image_text(self, img, page) -> str:
        """Run Tesseract OCR on a page image region and return text caption."""
        try:
            # Get image coordinates and extract
            bbox = (img['x0'], img['top'], img['x1'], img['bottom'])
            cropped_page = page.crop(bbox)
            
            # Convert to PIL Image
            img_obj = cropped_page.to_image()
            pil_image = img_obj.original
            
            # Use OCR to extract text
            text = pytesseract.image_to_string(pil_image)
            return text.strip()
            
        except Exception as e:
            logger.warning(f"OCR extraction failed: {str(e)}")
            return ""
    
    def _chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Document]:
        """Split text into sentence-aware chunks of ~N tokens with overlap."""
        if not text or not text.strip():
            return []
        
        # Tokenize into sentences
        sentences = sent_tokenize(text)
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = len(word_tokenize(sentence))
            
            # If adding this sentence would exceed chunk size, save current chunk
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append(Document(
                    content=chunk_text,
                    metadata=metadata.copy(),
                    content_type=metadata.get("type", "text")
                ))
                
                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(current_chunk)
                current_chunk = overlap_sentences + [sentence]
                current_tokens = sum(len(word_tokenize(s)) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        # Add final chunk if not empty
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(Document(
                content=chunk_text,
                metadata=metadata.copy(),
                content_type=metadata.get("type", "text")
            ))
        
        return chunks
    
    def _get_overlap_sentences(self, sentences: List[str]) -> List[str]:
        """Return trailing sentences to preserve context across consecutive chunks."""
        if len(sentences) <= 1:
            return []
        
        # Calculate how many sentences to include for overlap
        overlap_tokens = 0
        overlap_sentences = []
        
        for sentence in reversed(sentences):
            sentence_tokens = len(word_tokenize(sentence))
            if overlap_tokens + sentence_tokens <= self.chunk_overlap:
                overlap_sentences.insert(0, sentence)
                overlap_tokens += sentence_tokens
            else:
                break
        
        return overlap_sentences
